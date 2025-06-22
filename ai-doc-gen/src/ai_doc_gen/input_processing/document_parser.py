"""
Document Parser Module

Handles parsing of various document formats (PDF, DOCX, XML) and extracts
structured content for AI processing. Based on patterns from the original
codebase but enhanced for multi-format support.
"""

import logging
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any, Dict, List, Optional
import re

from pydantic import BaseModel

# PDF processing
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("pdfplumber not available - PDF parsing disabled")

# DOCX processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available - DOCX parsing disabled")

# XML processing
try:
    import xml.etree.ElementTree as ET
    XML_AVAILABLE = True
except ImportError:
    XML_AVAILABLE = False
    logging.warning("xml.etree not available - XML parsing disabled")

# HTML processing
try:
    import pandas as pd
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False
    logging.warning("beautifulsoup4 or pandas not available - HTML parsing disabled")

logger = logging.getLogger(__name__)

class ParsedDocument(BaseModel):
    """Structured representation of a parsed document."""
    filename: str
    file_type: str
    title: Optional[str] = None
    sections: List[Dict[str, Any]] = []
    metadata: Dict[str, Any] = {}
    raw_text: str = ""
    parsing_errors: List[str] = []

class DocumentParser(ABC):
    """Abstract base class for document parsers."""

    @abstractmethod
    def can_parse(self, file_path: str) -> bool:
        """Check if this parser can handle the given file."""
        pass

    @abstractmethod
    def parse(self, file_path: str) -> ParsedDocument:
        """Parse the document and return structured content."""
        pass

    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """Extract raw text from the document."""
        pass

class PDFParser(DocumentParser):
    """PDF document parser using pdfplumber."""

    def can_parse(self, file_path: str) -> bool:
        """Check if file is a PDF."""
        return (PDF_AVAILABLE and
                Path(file_path).suffix.lower() == '.pdf' and
                Path(file_path).exists())

    def parse(self, file_path: str) -> ParsedDocument:
        """Parse PDF document and extract structured content with performance optimization."""
        if not self.can_parse(file_path):
            raise ValueError(f"Cannot parse file: {file_path}")

        parsed_doc = ParsedDocument(
            filename=Path(file_path).name,
            file_type='pdf'
        )

        try:
            # Use caching for better performance
            cache_key = f"pdf_parse_{hash(file_path)}"
            
            # Check if we have cached results
            if hasattr(self, '_parse_cache') and cache_key in self._parse_cache:
                return self._parse_cache[cache_key]

            # Open PDF with optimized settings
            with pdfplumber.open(file_path) as pdf:
                
                # Extract title from metadata first
                if pdf.metadata and pdf.metadata.get('Title'):
                    parsed_doc.title = pdf.metadata['Title']
                
                # Process pages with optimized extraction
                sections = []
                raw_text_parts = []
                
                # Process pages in batches for better performance
                batch_size = 5
                for i in range(0, len(pdf.pages), batch_size):
                    batch_pages = pdf.pages[i:i + batch_size]
                    
                    for page_num, page in enumerate(batch_pages, i + 1):
                        try:
                            # Extract text with optimized settings
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                raw_text_parts.append(page_text)
                                
                                # Identify sections more efficiently
                                page_sections = self._identify_sections_optimized(page_text, page_num)
                                sections.extend(page_sections)
                                
                        except Exception as e:
                            error_msg = f"Error processing page {page_num}: {e}"
                            parsed_doc.parsing_errors.append(error_msg)
                            logger.warning(error_msg)

                # Combine results
                parsed_doc.sections = sections
                parsed_doc.raw_text = '\n'.join(raw_text_parts)
                
                # Add metadata
                parsed_doc.metadata = {
                    'page_count': len(pdf.pages),
                    'file_size': Path(file_path).stat().st_size,
                    'extraction_method': 'optimized_pdf_parser'
                }

                # Cache the results
                if not hasattr(self, '_parse_cache'):
                    self._parse_cache = {}
                self._parse_cache[cache_key] = parsed_doc

        except Exception as e:
            error_msg = f"Error parsing PDF {file_path}: {e}"
            parsed_doc.parsing_errors.append(error_msg)
            logger.error(error_msg)

        return parsed_doc

    def extract_text(self, file_path: str) -> str:
        """Extract raw text from PDF."""
        if not self.can_parse(file_path):
            raise ValueError(f"Cannot parse file: {file_path}")

        try:
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise

    def _identify_sections_optimized(self, page_text: str, page_num: int) -> List[Dict[str, Any]]:
        """Optimized section identification with better performance."""
        sections = []
        lines = page_text.split('\n')
        
        # Pre-compile regex patterns for better performance
        if not hasattr(self, '_header_patterns'):
            self._header_patterns = {
                'numbered': re.compile(r'^\d+\.\s+[A-Z]'),
                'roman': re.compile(r'^[IVX]+\.\s+[A-Z]'),
                'lettered': re.compile(r'^[A-Z]\.\s+[A-Z]'),
                'bold_caps': re.compile(r'^[A-Z][A-Z\s]{3,}$')
            }
        
        current_section = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if line is a header using optimized patterns
            if self._is_header_optimized(line):
                # Save previous section if exists
                if current_section:
                    current_section['content'] = current_content
                    sections.append(current_section)
                
                # Start new section
                current_section = {
                    'heading': line,
                    'level': self._get_header_level_optimized(line),
                    'content': [],
                    'page': page_num,
                    'start_line': len(sections)
                }
                current_content = []
            else:
                # Add to current content
                if current_section:
                    current_content.append(line)
        
        # Add final section
        if current_section:
            current_section['content'] = current_content
            sections.append(current_section)
        
        return sections

    def _is_header_optimized(self, line: str) -> bool:
        """Optimized header detection using pre-compiled patterns."""
        if len(line) < 3 or len(line) > 200:
            return False
            
        # Check for common header indicators
        if line.isupper() and len(line) > 5:
            return True
            
        # Use pre-compiled patterns
        for pattern in self._header_patterns.values():
            if pattern.match(line):
                return True
                
        # Check for technical section keywords
        tech_keywords = ['specification', 'installation', 'configuration', 'requirements', 
                        'overview', 'features', 'safety', 'maintenance', 'troubleshooting']
        if any(keyword in line.lower() for keyword in tech_keywords):
            return True
            
        return False

    def _get_header_level_optimized(self, header: str) -> int:
        """Optimized header level detection."""
        # Check for numbered patterns
        if re.match(r'^\d+\.', header):
            return 1
        elif re.match(r'^\d+\.\d+', header):
            return 2
        elif re.match(r'^\d+\.\d+\.\d+', header):
            return 3
        elif re.match(r'^[IVX]+\.', header):
            return 1
        elif re.match(r'^[A-Z]\.', header):
            return 2
        else:
            # Estimate level based on length and formatting
            if header.isupper() and len(header) > 10:
                return 1
            elif len(header) > 20:
                return 2
            else:
                return 3

class DOCXParser(DocumentParser):
    """DOCX document parser using python-docx."""

    def can_parse(self, file_path: str) -> bool:
        """Check if file is a DOCX."""
        return (DOCX_AVAILABLE and
                Path(file_path).suffix.lower() in ['.docx', '.doc'] and
                Path(file_path).exists())

    def parse(self, file_path: str) -> ParsedDocument:
        """Parse DOCX document and extract structured content."""
        if not self.can_parse(file_path):
            raise ValueError(f"Cannot parse file: {file_path}")

        filename = Path(file_path).name
        parsed_doc = ParsedDocument(
            filename=filename,
            file_type="docx",
            sections=[],
            metadata={},
            raw_text="",
            parsing_errors=[]
        )

        try:
            doc = Document(file_path)

            # Extract metadata
            core_props = doc.core_properties
            parsed_doc.metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "paragraphs": len(doc.paragraphs)
            }

            # Extract content
            sections = []
            current_section = None

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Check if this is a heading
                if self._is_heading(para):
                    # Save previous section
                    if current_section:
                        sections.append(current_section)

                    # Start new section
                    current_section = {
                        "heading": text,
                        "level": self._get_heading_level(para),
                        "content": [],
                        "page": 1  # DOCX doesn't have page info easily
                    }
                else:
                    # Add to current section
                    if current_section:
                        current_section["content"].append(text)
                    else:
                        # No section yet, might be title
                        if not parsed_doc.title:
                            parsed_doc.title = text

            # Add final section
            if current_section:
                sections.append(current_section)

            parsed_doc.sections = sections
            parsed_doc.raw_text = self.extract_text(file_path)

        except Exception as e:
            error_msg = f"Error parsing DOCX {file_path}: {e}"
            parsed_doc.parsing_errors.append(error_msg)
            logger.error(error_msg)

        return parsed_doc

    def extract_text(self, file_path: str) -> str:
        """Extract raw text from DOCX."""
        if not self.can_parse(file_path):
            raise ValueError(f"Cannot parse file: {file_path}")

        try:
            doc = Document(file_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise

    def _is_heading(self, paragraph) -> bool:
        """Check if paragraph is a heading."""
        # Check paragraph style
        if paragraph.style.name.startswith('Heading'):
            return True

        # Check if text looks like a heading
        text = paragraph.text.strip()
        if not text:
            return False

        # Simple heuristics
        if text.isupper() and len(text) < 100:
            return True

        if any(char.isdigit() for char in text) and len(text) < 50:
            return True

        return False

    def _get_heading_level(self, paragraph) -> int:
        """Get heading level from paragraph style."""
        style_name = paragraph.style.name
        if 'Heading' in style_name:
            # Extract number from style name
            import re
            match = re.search(r'Heading\s*(\d+)', style_name)
            if match:
                return int(match.group(1))

        # Default level based on text characteristics
        text = paragraph.text.strip()
        if text.isupper():
            return 1
        elif any(char.isdigit() for char in text):
            return 2
        else:
            return 3

class XMLParser(DocumentParser):
    """XML document parser using ElementTree."""

    def can_parse(self, file_path: str) -> bool:
        """Check if file is an XML."""
        return (XML_AVAILABLE and
                Path(file_path).suffix.lower() == '.xml' and
                Path(file_path).exists())

    def parse(self, file_path: str) -> ParsedDocument:
        """Parse XML document and extract structured content."""
        if not self.can_parse(file_path):
            raise ValueError(f"Cannot parse file: {file_path}")

        filename = Path(file_path).name
        parsed_doc = ParsedDocument(
            filename=filename,
            file_type="xml",
            sections=[],
            metadata={},
            raw_text="",
            parsing_errors=[]
        )

        try:
            tree = ET.parse(file_path)
            root = tree.getroot()

            # Extract metadata from root attributes
            parsed_doc.metadata = dict(root.attrib)

            # Extract title
            title_elem = root.find('.//title')
            if title_elem is not None:
                parsed_doc.title = title_elem.text

            # Extract sections
            sections = []
            for section_elem in root.findall('.//section'):
                section = self._parse_section_element(section_elem)
                if section:
                    sections.append(section)

            parsed_doc.sections = sections
            parsed_doc.raw_text = self.extract_text(file_path)

        except Exception as e:
            error_msg = f"Error parsing XML {file_path}: {e}"
            parsed_doc.parsing_errors.append(error_msg)
            logger.error(error_msg)

        return parsed_doc

    def extract_text(self, file_path: str) -> str:
        """Extract raw text from XML."""
        if not self.can_parse(file_path):
            raise ValueError(f"Cannot parse file: {file_path}")

        try:
            tree = ET.parse(file_path)
            root = tree.getroot()

            # Extract all text content
            text_parts = []
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    text_parts.append(elem.text.strip())

            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from XML {file_path}: {e}")
            raise

    def _parse_section_element(self, section_elem) -> Optional[Dict[str, Any]]:
        """Parse a section element and extract structured content."""
        try:
            # Extract heading
            heading_elem = section_elem.find('heading')
            if heading_elem is None:
                return None

            heading = heading_elem.text.strip() if heading_elem.text else ""

            # Extract content
            content_elem = section_elem.find('content')
            content = []
            if content_elem is not None and content_elem.text:
                content = [line.strip() for line in content_elem.text.split('\n') if line.strip()]

            return {
                "heading": heading,
                "level": 1,  # Default level for XML sections
                "content": content,
                "source": "xml"
            }
        except Exception as e:
            logger.warning(f"Error parsing XML section: {e}")
            return None

class HTMLParser(DocumentParser):
    """HTML document parser using BeautifulSoup."""

    def can_parse(self, file_path: str) -> bool:
        """Check if file is an HTML file."""
        return (HTML_AVAILABLE and
                Path(file_path).suffix.lower() in ['.html', '.htm'] and
                Path(file_path).exists())

    def parse(self, file_path: str) -> ParsedDocument:
        """Parse HTML document and extract structured content using enhanced scraping methods with quality assessment."""
        if not self.can_parse(file_path):
            raise ValueError(f"Cannot parse file: {file_path}")

        filename = Path(file_path).name
        parsed_doc = ParsedDocument(
            filename=filename,
            file_type="html",
            sections=[],
            metadata={},
            raw_text="",
            parsing_errors=[]
        )

        try:
            with open(file_path, encoding='utf-8') as f:
                html_content = f.read()

            soup = BeautifulSoup(html_content, 'html.parser')

            # Extract metadata
            parsed_doc.metadata = {
                "title": soup.title.string if soup.title else "",
                "language": soup.get('lang', ''),
                "charset": soup.meta.get('charset', '') if soup.meta else "",
                "description": soup.find('meta', {'name': 'description'}).get('content', '') if soup.find('meta', {'name': 'description'}) else ""
            }

            # Extract title
            if soup.title:
                parsed_doc.title = soup.title.string.strip()

            # Extract sections using enhanced methods with quality assessment
            sections = self._extract_enhanced_sections(soup)
            parsed_doc.sections = sections

            # Extract raw text
            parsed_doc.raw_text = self._extract_clean_text(soup)

        except Exception as e:
            error_msg = f"Error parsing HTML {file_path}: {e}"
            parsed_doc.parsing_errors.append(error_msg)
            logger.error(error_msg)

        return parsed_doc

    def extract_text(self, file_path: str) -> str:
        """Extract raw text from HTML."""
        if not self.can_parse(file_path):
            raise ValueError(f"Cannot parse file: {file_path}")

        try:
            with open(file_path, encoding='utf-8') as f:
                html_content = f.read()

            soup = BeautifulSoup(html_content, 'html.parser')
            return self._extract_clean_text(soup)
        except Exception as e:
            logger.error(f"Error extracting text from HTML {file_path}: {e}")
            raise

    def _extract_enhanced_sections(self, soup) -> List[Dict[str, Any]]:
        """Extract sections from HTML using enhanced scraping methods with quality assessment."""
        sections = []

        # First, try to find main content area
        main_content = self._find_main_content(soup)
        if main_content:
            soup = main_content

        # Find all heading elements (h1, h2, h3, h4, h5, h6)
        headings = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])

        for heading in headings:
            try:
                heading_text = heading.get_text(strip=True)
                if not heading_text or len(heading_text) < 2:
                    continue

                # Determine heading level
                level = int(heading.name[1])

                # Extract content following this heading using enhanced methods
                content = self._extract_enhanced_section_content(heading)
                
                # Quality assessment for this section
                quality_score = self._assess_section_quality(heading_text, content)
                
                # Only include sections with reasonable quality
                if quality_score > 0.3:  # Minimum quality threshold
                    section = {
                        "heading": heading_text,
                        "level": level,
                        "content": content,
                        "source": "html_enhanced",
                        "quality_score": quality_score,
                        "content_length": sum(len(c) for c in content),
                        "content_preview": self._create_content_preview(content)
                    }
                    sections.append(section)

            except Exception as e:
                logger.warning(f"Error extracting section from HTML: {e}")

        # If no headings found or quality is poor, try alternative extraction methods
        if not sections or len(sections) < 3:
            alternative_sections = self._extract_alternative_sections(soup)
            sections.extend(alternative_sections)

        # Sort sections by quality score and remove duplicates
        sections = self._deduplicate_and_sort_sections(sections)

        return sections

    def _find_main_content(self, soup) -> Optional[BeautifulSoup]:
        """Find the main content area of the HTML document."""
        # Common selectors for main content
        main_selectors = [
            'main',
            '[role="main"]',
            '.main-content',
            '.content',
            '#content',
            '#main',
            '.post-content',
            '.article-content',
            '.document-content'
        ]
        
        for selector in main_selectors:
            main_elem = soup.select_one(selector)
            if main_elem and len(main_elem.get_text(strip=True)) > 500:
                return main_elem
        
        # If no main content found, try to identify the largest content area
        content_areas = soup.find_all(['div', 'section', 'article'])
        if content_areas:
            largest_area = max(content_areas, key=lambda x: len(x.get_text(strip=True)))
            if len(largest_area.get_text(strip=True)) > 1000:
                return largest_area
        
        return None

    def _assess_section_quality(self, heading: str, content: List[str]) -> float:
        """Assess the quality of a section based on heading and content."""
        score = 0.0
        
        # Heading quality (30% of score)
        if len(heading) > 5 and len(heading) < 200:
            score += 0.3
        
        # Content quality (70% of score)
        if content:
            total_content_length = sum(len(c) for c in content)
            if total_content_length > 50:  # Minimum content length
                score += 0.2
            
            # Check for meaningful content (not just whitespace or single words)
            meaningful_content = [c for c in content if len(c.strip()) > 10]
            if len(meaningful_content) > 0:
                score += 0.3
            
            # Check for technical content indicators
            technical_indicators = ['specification', 'installation', 'configuration', 'requirements', 
                                  'features', 'overview', 'description', 'procedure', 'warning']
            if any(indicator in heading.lower() for indicator in technical_indicators):
                score += 0.2
        
        return min(score, 1.0)

    def _extract_alternative_sections(self, soup) -> List[Dict[str, Any]]:
        """Extract sections using alternative methods when headings are not available."""
        sections = []
        
        # Method 1: Extract from structural elements with meaningful content
        structural_elements = soup.find_all(['div', 'section', 'article'])
        
        for elem in structural_elements:
            text = elem.get_text(strip=True)
            if len(text) < 100:  # Skip very short elements
                continue
            
            # Try to find a meaningful heading
            heading = self._extract_heading_from_element(elem)
            if not heading:
                continue
            
            content = [p.get_text(strip=True) for p in elem.find_all(['p', 'li', 'td']) 
                      if p.get_text(strip=True) and len(p.get_text(strip=True)) > 20]
            
            if content:
                quality_score = self._assess_section_quality(heading, content)
                if quality_score > 0.4:
                    section = {
                        "heading": heading,
                        "level": 1,
                        "content": content,
                        "source": "html_alternative",
                        "quality_score": quality_score,
                        "content_length": sum(len(c) for c in content),
                        "content_preview": self._create_content_preview(content)
                    }
                    sections.append(section)
        
        return sections

    def _extract_heading_from_element(self, elem) -> Optional[str]:
        """Extract a meaningful heading from an element."""
        # Try to find existing headings
        heading = elem.find(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
        if heading:
            return heading.get_text(strip=True)
        
        # Try to extract from element attributes
        if elem.get('id'):
            return elem.get('id').replace('-', ' ').replace('_', ' ').title()
        
        if elem.get('class'):
            class_name = elem.get('class')[0] if elem.get('class') else ''
            if class_name and len(class_name) > 3:
                return class_name.replace('-', ' ').replace('_', ' ').title()
        
        # Try to extract from first paragraph or list item
        first_content = elem.find(['p', 'li', 'strong', 'b'])
        if first_content:
            text = first_content.get_text(strip=True)
            if len(text) > 10 and len(text) < 100:
                return text
        
        return None

    def _create_content_preview(self, content: List[str]) -> str:
        """Create a preview of the content for display."""
        if not content:
            return ""
        
        # Take first meaningful content item
        for item in content:
            if len(item.strip()) > 20:
                preview = item.strip()[:100]
                return preview + "..." if len(item) > 100 else preview
        
        return content[0][:100] + "..." if len(content[0]) > 100 else content[0]

    def _deduplicate_and_sort_sections(self, sections: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Remove duplicate sections and sort by quality."""
        # Remove duplicates based on heading similarity
        unique_sections = []
        seen_headings = set()
        
        for section in sections:
            heading_lower = section["heading"].lower().strip()
            if heading_lower not in seen_headings:
                seen_headings.add(heading_lower)
                unique_sections.append(section)
        
        # Sort by quality score (highest first)
        unique_sections.sort(key=lambda x: x.get("quality_score", 0), reverse=True)
        
        return unique_sections

    def _extract_enhanced_section_content(self, heading) -> List[str]:
        """Extract content that follows a heading using enhanced methods including table extraction."""
        content = []
        current = heading.next_sibling

        while current:
            if current.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                break  # Stop at next heading

            if hasattr(current, 'get_text'):
                # Extract text content
                text = current.get_text(strip=True)
                if text:
                    content.append(text)

                # Extract tables using pandas if available
                if current.name == 'table':
                    try:
                        table_content = self._extract_table_content(current)
                        if table_content:
                            content.extend(table_content)
                    except Exception as e:
                        logger.warning(f"Error extracting table: {e}")

            current = current.next_sibling

        return content

    def _extract_table_content(self, table_element) -> List[str]:
        """Extract structured content from HTML tables using pandas."""
        try:
            # Use pandas to read HTML table
            table_html = str(table_element)
            tables = pd.read_html(table_html)

            table_content = []
            for i, df in enumerate(tables):
                if not df.empty:
                    # Convert DataFrame to readable text
                    table_text = f"Table {i+1}:\n"
                    table_text += df.to_string(index=False)
                    table_content.append(table_text)

            return table_content
        except Exception as e:
            logger.warning(f"Error extracting table with pandas: {e}")
            # Fallback to BeautifulSoup extraction
            return self._extract_table_with_bs4(table_element)

    def _extract_table_with_bs4(self, table_element) -> List[str]:
        """Fallback table extraction using BeautifulSoup."""
        table_content = []

        # Extract table headers
        headers = []
        for th in table_element.find_all('th'):
            headers.append(th.get_text(strip=True))

        if headers:
            table_content.append("Headers: " + " | ".join(headers))

        # Extract table rows
        for tr in table_element.find_all('tr'):
            cells = []
            for td in tr.find_all(['td', 'th']):
                cells.append(td.get_text(strip=True))
            if cells:
                table_content.append(" | ".join(cells))

        return table_content

    def _extract_clean_text(self, soup) -> str:
        """Extract clean text from HTML, removing scripts, styles, and navigation."""
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "header", "footer"]):
            script.decompose()

        # Get text and clean it up
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)

        return text

class TextParser(DocumentParser):
    """Plain text document parser."""

    def can_parse(self, file_path: str) -> bool:
        """Check if file is a text file."""
        return (Path(file_path).suffix.lower() == '.txt' and
                Path(file_path).exists())

    def parse(self, file_path: str) -> ParsedDocument:
        """Parse text document and extract structured content."""
        if not self.can_parse(file_path):
            raise ValueError(f"Cannot parse file: {file_path}")

        filename = Path(file_path).name
        parsed_doc = ParsedDocument(
            filename=filename,
            file_type="txt",
            sections=[],
            metadata={},
            raw_text="",
            parsing_errors=[]
        )

        try:
            with open(file_path, encoding='utf-8') as f:
                content = f.read()

            parsed_doc.raw_text = content

            # Extract title from first line
            lines = content.split('\n')
            if lines:
                first_line = lines[0].strip()
                if first_line.startswith('# '):
                    parsed_doc.title = first_line[2:]  # Remove '# ' prefix
                elif first_line:
                    parsed_doc.title = first_line

            # Extract sections based on markdown headers
            sections = []
            current_section = None
            current_content = []

            for line_num, line in enumerate(lines, 1):
                line = line.strip()
                if not line:
                    continue

                # Check for headers (# ## ###)
                if line.startswith('#'):
                    # Save previous section
                    if current_section:
                        current_section['content'] = current_content
                        sections.append(current_section)

                    # Start new section
                    level = len(line) - len(line.lstrip('#'))
                    title = line.lstrip('#').strip()
                    current_section = {
                        'heading': title,
                        'level': level,
                        'content': [],
                        'page': 1  # Text files don't have pages
                    }
                    current_content = []
                else:
                    # Add to current section content
                    if current_section:
                        current_content.append(line)
                    else:
                        # No section yet, might be title
                        if not parsed_doc.title:
                            parsed_doc.title = line

            # Add final section
            if current_section:
                current_section['content'] = current_content
                sections.append(current_section)

            parsed_doc.sections = sections

        except Exception as e:
            error_msg = f"Error parsing text file {file_path}: {e}"
            parsed_doc.parsing_errors.append(error_msg)
            logger.error(error_msg)

        return parsed_doc

    def extract_text(self, file_path: str) -> str:
        """Extract raw text from text file."""
        if not self.can_parse(file_path):
            raise ValueError(f"Cannot parse file: {file_path}")

        try:
            with open(file_path, encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise

class DocumentParserFactory:
    """Factory for creating appropriate document parsers."""

    def __init__(self):
        """Initialize available parsers."""
        self.parsers = []

        if PDF_AVAILABLE:
            self.parsers.append(PDFParser())
        if DOCX_AVAILABLE:
            self.parsers.append(DOCXParser())
        if XML_AVAILABLE:
            self.parsers.append(XMLParser())
        if HTML_AVAILABLE:
            self.parsers.append(HTMLParser())

        # Always add text parser
        self.parsers.append(TextParser())

    def get_parser(self, file_path: str) -> Optional[DocumentParser]:
        """Get appropriate parser for the given file."""
        for parser in self.parsers:
            if parser.can_parse(file_path):
                return parser
        return None

    def parse_document(self, file_path: str) -> ParsedDocument:
        """Parse document using appropriate parser."""
        parser = self.get_parser(file_path)
        if not parser:
            raise ValueError(f"No parser available for file: {file_path}")

        return parser.parse(file_path)

    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        formats = []
        for parser in self.parsers:
            if isinstance(parser, PDFParser):
                formats.append("PDF")
            elif isinstance(parser, DOCXParser):
                formats.append("DOCX")
            elif isinstance(parser, XMLParser):
                formats.append("XML")
            elif isinstance(parser, HTMLParser):
                formats.append("HTML")
            elif isinstance(parser, TextParser):
                formats.append("TXT")
        return formats

# Convenience function
def parse_document(file_path: str) -> ParsedDocument:
    """Parse a document using the appropriate parser."""
    factory = DocumentParserFactory()
    return factory.parse_document(file_path)
